from __future__ import annotations

import copy
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


TEMPLATE_PATH = Path("/Users/caobingbing/Desktop/接口文档.docx")
OUTPUT_PATH = Path("/Users/caobingbing/Desktop/接口文档-系统补全版.docx")


def cn_num(value: int) -> str:
    digits = "零一二三四五六七八九"
    if value <= 10:
        return "十" if value == 10 else digits[value]
    if value < 20:
        return "十" + digits[value % 10]
    tens, ones = divmod(value, 10)
    return digits[tens] + "十" + (digits[ones] if ones else "")


def clear_document(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    doc.add_paragraph(text, style=style)


def add_lines(doc: Document, lines: list[str], style: str = "Normal") -> None:
    for line in lines:
        add_paragraph(doc, line, style)


def add_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Normal Table"
    table.autofit = False
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value
            if widths and idx < len(widths):
                row[idx].width = Inches(widths[idx])


def dumps(data) -> list[str]:
    return json.dumps(data, ensure_ascii=False, indent=2).splitlines()


def param_row(position: str, name: str, type_: str, required: str, desc: str) -> list[str]:
    return [position, name, type_, required, desc]


def field_row(name: str, type_: str, desc: str) -> list[str]:
    return [name, type_, desc]


COMMON_RESPONSE_FIELDS = [
    field_row("code", "number", "业务状态码，0=成功，其他值=失败"),
    field_row("message", "string", "接口提示信息，成功时通常为 success"),
]


def response_rows(data_rows: list[list[str]]) -> list[list[str]]:
    return [["字段名", "数据类型", "字段说明"], *COMMON_RESPONSE_FIELDS, field_row("data", "object/null", "业务数据主体，具体结构见下方"), *data_rows]


def void_response_rows() -> list[list[str]]:
    return [["字段名", "数据类型", "字段说明"], *COMMON_RESPONSE_FIELDS, field_row("data", "null", "无业务数据，固定为 null")]


AUTH_LOGIN_FIELDS = [
    field_row("data.token", "string", "JWT 登录令牌，后续放入 Authorization: Bearer {token}"),
    field_row("data.user.userId", "number", "用户 ID"),
    field_row("data.user.phone", "string", "手机号"),
    field_row("data.user.nickname", "string", "昵称"),
    field_row("data.user.avatarUrl", "string", "头像地址"),
    field_row("data.user.role", "string", "角色，USER 或 ADMIN"),
    field_row("data.user.realNameStatus", "string", "实名状态，UNSUBMITTED/PENDING/APPROVED/REJECTED"),
    field_row("data.user.gender", "string", "性别"),
    field_row("data.user.contactPhone", "string", "联系手机号"),
    field_row("data.user.bio", "string", "个人简介"),
]

USER_PROFILE_FIELDS = [
    field_row("data.userId", "number", "用户 ID"),
    field_row("data.phone", "string", "登录手机号"),
    field_row("data.nickname", "string", "昵称"),
    field_row("data.avatarUrl", "string", "头像地址"),
    field_row("data.role", "string", "角色"),
    field_row("data.realNameStatus", "string", "实名状态"),
    field_row("data.gender", "string", "性别"),
    field_row("data.contactPhone", "string", "联系手机号"),
    field_row("data.bio", "string", "个人简介"),
]

VERIFICATION_FIELDS = [
    field_row("data.id", "number", "实名记录 ID"),
    field_row("data.realName", "string", "真实姓名"),
    field_row("data.idCardNo", "string", "身份证号"),
    field_row("data.frontFileId", "number", "身份证正面文件 ID"),
    field_row("data.frontFileUrl", "string", "身份证正面访问地址"),
    field_row("data.backFileId", "number", "身份证反面文件 ID"),
    field_row("data.backFileUrl", "string", "身份证反面访问地址"),
    field_row("data.status", "string", "审核状态，PENDING/APPROVED/REJECTED"),
    field_row("data.rejectReason", "string", "驳回原因"),
    field_row("data.reviewedBy", "number", "审核管理员 ID"),
    field_row("data.reviewedAt", "string(datetime)", "审核时间"),
]

ERRAND_RULE_FIELDS = [
    field_row("data.urgentFee", "number", "加急附加费"),
    field_row("data.fragileFee", "number", "易碎附加费"),
]

FILE_FIELDS = [
    field_row("fileId", "number", "文件 ID"),
    field_row("url", "string", "文件访问地址"),
    field_row("originalName", "string", "原始文件名"),
]

REPORT_FIELDS = [
    field_row("data.id", "number", "举报单 ID"),
    field_row("data.module", "string", "业务模块，如 errand"),
    field_row("data.targetType", "string", "目标类型，如 order"),
    field_row("data.targetId", "number", "被举报目标 ID"),
    field_row("data.reportType", "string", "举报类型"),
    field_row("data.description", "string", "举报说明"),
    field_row("data.contactPhone", "string", "联系电话"),
    field_row("data.status", "string", "状态，PENDING/PROCESSING/RESOLVED/REJECTED"),
    field_row("data.handleRemark", "string", "处理备注"),
    field_row("data.attachments[].fileId", "number", "附件文件 ID"),
    field_row("data.attachments[].url", "string", "附件访问地址"),
    field_row("data.attachments[].originalName", "string", "附件原始名称"),
    field_row("data.handledBy", "number", "处理人管理员 ID"),
    field_row("data.handledAt", "string(datetime)", "处理时间"),
    field_row("data.createdAt", "string(datetime)", "创建时间"),
]

MESSAGE_FIELDS = [
    field_row("data.id", "number", "消息 ID"),
    field_row("data.conversationId", "number", "所属会话 ID"),
    field_row("data.senderId", "number", "发送人 ID，系统通知可能为 0"),
    field_row("data.type", "string", "消息类型，TEXT/IMAGE"),
    field_row("data.content", "string", "消息内容"),
    field_row("data.createdAt", "string(datetime)", "发送时间"),
]

CONVERSATION_FIELDS = [
    field_row("data[].id", "number", "会话 ID"),
    field_row("data[].type", "string", "会话类型，PRIVATE/SYSTEM"),
    field_row("data[].title", "string", "会话标题"),
    field_row("data[].unreadCount", "number", "当前会话未读数"),
    field_row("data[].lastMessage.id", "number", "最后一条消息 ID"),
    field_row("data[].lastMessage.type", "string", "最后一条消息类型"),
    field_row("data[].lastMessage.content", "string", "最后一条消息内容"),
    field_row("data[].lastMessage.createdAt", "string(datetime)", "最后一条消息时间"),
    field_row("data[].updatedAt", "string(datetime)", "会话更新时间"),
]

MESSAGES_FIELDS = [
    field_row("data[].id", "number", "消息 ID"),
    field_row("data[].conversationId", "number", "所属会话 ID"),
    field_row("data[].senderId", "number", "发送人 ID"),
    field_row("data[].type", "string", "消息类型"),
    field_row("data[].content", "string", "消息内容"),
    field_row("data[].createdAt", "string(datetime)", "发送时间"),
]

UNREAD_FIELDS = [
    field_row("data.totalUnread", "number", "全部会话未读总数"),
]

UPLOAD_FIELDS = [
    field_row("data.fileId", "number", "文件 ID"),
    field_row("data.url", "string", "图片访问地址"),
    field_row("data.originalName", "string", "原始文件名"),
]

ERRAND_CONVERSATION_FIELDS = [
    field_row("data.conversationId", "number", "订单对应的私聊会话 ID"),
]

ADMIN_USER_DETAIL_FIELDS = [
    field_row("data.userId", "number", "用户 ID"),
    field_row("data.phone", "string", "手机号"),
    field_row("data.nickname", "string", "昵称"),
    field_row("data.avatarUrl", "string", "头像地址"),
    field_row("data.role", "string", "角色"),
    field_row("data.status", "string", "账号状态，ACTIVE/DISABLED"),
    field_row("data.realNameStatus", "string", "实名状态"),
    field_row("data.gender", "string", "性别"),
    field_row("data.contactPhone", "string", "联系手机号"),
    field_row("data.bio", "string", "个人简介"),
    field_row("data.reportCount", "number", "该用户举报次数统计"),
    field_row("data.createdAt", "string(datetime)", "注册时间"),
]

ADMIN_USER_LIST_FIELDS = [
    field_row("data[].id", "number", "用户 ID"),
    field_row("data[].phone", "string", "手机号"),
    field_row("data[].nickname", "string", "昵称"),
    field_row("data[].avatarUrl", "string", "头像地址"),
    field_row("data[].role", "string", "角色"),
    field_row("data[].status", "string", "账号状态"),
    field_row("data[].createdAt", "string(datetime)", "创建时间"),
]

ADMIN_VERIFICATION_PAGE_FIELDS = [
    field_row("data.current", "number", "当前页码"),
    field_row("data.size", "number", "每页条数"),
    field_row("data.total", "number", "总记录数"),
    field_row("data.pages", "number", "总页数"),
    field_row("data.records[].id", "number", "实名记录 ID"),
    field_row("data.records[].userId", "number", "用户 ID"),
    field_row("data.records[].realName", "string", "真实姓名"),
    field_row("data.records[].idCardNo", "string", "身份证号"),
    field_row("data.records[].frontFileId", "number", "身份证正面文件 ID"),
    field_row("data.records[].backFileId", "number", "身份证反面文件 ID"),
    field_row("data.records[].status", "string", "审核状态"),
    field_row("data.records[].rejectReason", "string", "驳回原因"),
    field_row("data.records[].reviewedBy", "number", "审核管理员 ID"),
    field_row("data.records[].reviewedAt", "string(datetime)", "审核时间"),
    field_row("data.records[].createdAt", "string(datetime)", "提交时间"),
]

ADMIN_REPORT_PAGE_FIELDS = [
    field_row("data.current", "number", "当前页码"),
    field_row("data.size", "number", "每页条数"),
    field_row("data.total", "number", "总记录数"),
    field_row("data.pages", "number", "总页数"),
    field_row("data.records[].id", "number", "举报单 ID"),
    field_row("data.records[].module", "string", "业务模块"),
    field_row("data.records[].targetType", "string", "目标类型"),
    field_row("data.records[].targetId", "number", "目标 ID"),
    field_row("data.records[].reportType", "string", "举报类型"),
    field_row("data.records[].description", "string", "举报说明"),
    field_row("data.records[].contactPhone", "string", "联系电话"),
    field_row("data.records[].status", "string", "处理状态"),
    field_row("data.records[].submitterId", "number", "提交人用户 ID"),
    field_row("data.records[].handleRemark", "string", "处理备注"),
    field_row("data.records[].handledBy", "number", "处理管理员 ID"),
    field_row("data.records[].handledAt", "string(datetime)", "处理时间"),
    field_row("data.records[].createdAt", "string(datetime)", "创建时间"),
]

ERRAND_ORDER_FIELDS = [
    field_row("data.id", "number", "订单 ID"),
    field_row("data.orderNo", "string", "订单编号"),
    field_row("data.serviceType", "string", "服务类型，PICKUP/DELIVERY/PURCHASE/PRINT"),
    field_row("data.pickupAddress", "string", "取货地址"),
    field_row("data.deliveryAddress", "string", "送达地址"),
    field_row("data.pickupTimeText", "string", "期望取货时间说明"),
    field_row("data.detailContent", "string", "订单详细需求"),
    field_row("data.remark", "string", "补充备注"),
    field_row("data.baseFee", "number", "基础费用"),
    field_row("data.urgentFee", "number", "加急费"),
    field_row("data.fragileFee", "number", "易碎费"),
    field_row("data.totalFee", "number", "订单总费用"),
    field_row("data.urgentFlag", "boolean", "是否加急"),
    field_row("data.fragileFlag", "boolean", "是否易碎"),
    field_row("data.status", "string", "订单状态"),
    field_row("data.acceptDeadline", "string(datetime)", "接单截止时间"),
    field_row("data.acceptedAt", "string(datetime)", "接单时间"),
    field_row("data.pickedUpAt", "string(datetime)", "取货时间"),
    field_row("data.deliveredAt", "string(datetime)", "送达/处理中时间"),
    field_row("data.completedAt", "string(datetime)", "完成时间"),
    field_row("data.cancelledAt", "string(datetime)", "取消时间"),
    field_row("data.cancelReason", "string", "取消原因"),
    field_row("data.publicVisible", "boolean", "是否仍在广场公开展示"),
    field_row("data.publisher.userId", "number", "发布人用户 ID"),
    field_row("data.publisher.nickname", "string", "发布人昵称"),
    field_row("data.publisher.phone", "string", "发布人手机号"),
    field_row("data.runner.userId", "number", "接单人用户 ID"),
    field_row("data.runner.nickname", "string", "接单人昵称"),
    field_row("data.runner.phone", "string", "接单人手机号"),
    field_row("data.attachments[].fileId", "number", "附件文件 ID"),
    field_row("data.attachments[].url", "string", "附件访问地址"),
    field_row("data.attachments[].originalName", "string", "附件原文件名"),
    field_row("data.relatedReports[].id", "number", "关联举报单 ID"),
    field_row("data.relatedReports[].status", "string", "举报处理状态"),
    field_row("data.conversationId", "number", "关联会话 ID"),
    field_row("data.canAccept", "boolean", "当前用户是否可接单"),
    field_row("data.canCancel", "boolean", "当前用户是否可取消"),
    field_row("data.canStart", "boolean", "当前用户是否可开始执行"),
    field_row("data.canDeliver", "boolean", "当前用户是否可提交送达"),
    field_row("data.canComplete", "boolean", "当前用户是否可确认完成"),
    field_row("data.canReport", "boolean", "当前用户是否可发起举报"),
    field_row("data.createdAt", "string(datetime)", "创建时间"),
]

ERRAND_ORDER_LIST_FIELDS = [
    field_row("data[].id", "number", "订单 ID"),
    field_row("data[].orderNo", "string", "订单编号"),
    field_row("data[].serviceType", "string", "服务类型"),
    field_row("data[].pickupAddress", "string", "取货地址"),
    field_row("data[].deliveryAddress", "string", "送达地址"),
    field_row("data[].totalFee", "number", "订单总费用"),
    field_row("data[].status", "string", "订单状态"),
    field_row("data[].publisher.nickname", "string", "发布人昵称"),
    field_row("data[].runner.nickname", "string", "接单人昵称"),
    field_row("data[].conversationId", "number", "关联会话 ID"),
    field_row("data[].createdAt", "string(datetime)", "创建时间"),
]

ERRAND_ORDER_PAGE_FIELDS = [
    field_row("data.current", "number", "当前页码"),
    field_row("data.size", "number", "每页条数"),
    field_row("data.total", "number", "总记录数"),
    field_row("data.pages", "number", "总页数"),
    field_row("data.records[].id", "number", "订单 ID"),
    field_row("data.records[].orderNo", "string", "订单编号"),
    field_row("data.records[].serviceType", "string", "服务类型"),
    field_row("data.records[].pickupAddress", "string", "取货地址"),
    field_row("data.records[].deliveryAddress", "string", "送达地址"),
    field_row("data.records[].totalFee", "number", "订单总费用"),
    field_row("data.records[].status", "string", "订单状态"),
    field_row("data.records[].publisher.nickname", "string", "发布人昵称"),
    field_row("data.records[].runner.nickname", "string", "接单人昵称"),
    field_row("data.records[].conversationId", "number", "关联会话 ID"),
    field_row("data.records[].createdAt", "string(datetime)", "创建时间"),
]


def success_lines(payload) -> list[str]:
    return dumps(payload)


def fail_lines(code: int, message: str) -> list[str]:
    return dumps({"code": code, "message": message, "data": None})


def endpoint(
    title: str,
    category: str,
    name: str,
    path: str,
    method: str,
    content_type: str,
    feature: str,
    purpose: str,
    request_rows: list[list[str]],
    request_example: list[str],
    response_data_rows: list[list[str]],
    success_example: list[str],
    fail_examples: list[tuple[str, list[str]]],
    logic_lines: list[str],
    section_header: str | None = None,
) -> dict:
    return {
        "title": title,
        "category": category,
        "name": name,
        "path": path,
        "method": method,
        "content_type": content_type,
        "feature": feature,
        "purpose": purpose,
        "request_rows": request_rows,
        "request_example": request_example,
        "response_rows": response_rows(response_data_rows) if response_data_rows else void_response_rows(),
        "success_example": success_example,
        "fail_examples": fail_examples,
        "logic_lines": logic_lines,
        "section_header": section_header,
    }


USER_TOKEN = "Bearer {token}"
ADMIN_TOKEN = "Bearer {adminToken}"


endpoints = [
    endpoint(
        "用户注册接口",
        "认证接口",
        "用户注册",
        "/api/auth/register",
        "POST",
        "application/json",
        "无需登录即可访问",
        "注册普通用户账号，校验手机号和密码后自动返回登录令牌与用户信息。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Body", "phone", "string", "是", "注册手机号，必须匹配 1 开头 11 位手机号"),
            param_row("Body", "password", "string", "是", "登录密码，长度 6-16 位"),
            param_row("Body", "confirmPassword", "string", "是", "确认密码，需与 password 一致"),
        ],
        success_lines({"phone": "13800138000", "password": "123456", "confirmPassword": "123456"}),
        AUTH_LOGIN_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "token": "eyJhbGciOiJIUzI1NiJ9.register",
                    "user": {
                        "userId": 1001,
                        "phone": "13800138000",
                        "nickname": "用户8000",
                        "avatarUrl": None,
                        "role": "USER",
                        "realNameStatus": "UNSUBMITTED",
                        "gender": None,
                        "contactPhone": "13800138000",
                        "bio": None,
                    },
                },
            }
        ),
        [
            ("失败返回（两次密码不一致）", fail_lines(400, "两次输入密码不一致")),
            ("失败返回（手机号已注册）", fail_lines(400, "手机号已注册")),
        ],
        [
            "1. 该接口属于账号创建接口，使用 POST 以承载表单数据并避免敏感信息出现在 URL 中。",
            "2. 后端会先校验确认密码与手机号唯一性，再对密码做 BCrypt 加密后写入用户表与用户资料表。",
            "3. 注册成功后系统会直接签发 JWT，并返回初始化后的用户资料，前端可跳过二次登录流程。",
        ],
    ),
    endpoint(
        "用户登录接口",
        "核心鉴权接口",
        "用户登录",
        "/api/auth/login",
        "POST",
        "application/json",
        "无需登录即可访问",
        "使用手机号和密码完成身份认证，成功后返回 JWT 令牌与当前用户资料。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Body", "phone", "string", "是", "登录手机号，格式必须正确"),
            param_row("Body", "password", "string", "是", "登录密码"),
        ],
        success_lines({"phone": "13800138000", "password": "123456"}),
        AUTH_LOGIN_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "token": "eyJhbGciOiJIUzI1NiJ9.login",
                    "user": {
                        "userId": 1001,
                        "phone": "13800138000",
                        "nickname": "小陈",
                        "avatarUrl": "http://localhost:9001/api/files/view/avatar.png",
                        "role": "USER",
                        "realNameStatus": "APPROVED",
                        "gender": "男",
                        "contactPhone": "13800138000",
                        "bio": "校园跑腿达人",
                    },
                },
            }
        ),
        [("失败返回（账号或密码错误）", fail_lines(401, "手机号或密码错误"))],
        [
            "1. 登录接口是移动端与后台管理端获取令牌的统一入口，普通用户与管理员都走同一套认证逻辑。",
            "2. 后端根据手机号查询用户，并使用 BCrypt 校验密码，校验失败返回 code=401。",
            "3. 前端接收 token 后应保存到本地，并在后续请求头中统一携带 Authorization: Bearer {token}。",
        ],
    ),
    endpoint(
        "用户退出登录接口",
        "会话管理接口",
        "用户退出登录",
        "/api/auth/logout",
        "POST",
        "application/json",
        "需携带普通用户 Token 或管理员 Token",
        "通知前端完成本地登录态清理，当前版本后端不做服务端会话持久化注销。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"登录令牌，格式 {USER_TOKEN}"),
        ],
        ["POST /api/auth/logout", f"Authorization: {USER_TOKEN}", "", "{}"],
        [],
        success_lines({"code": 0, "message": "退出成功", "data": None}),
        [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
        [
            "1. 该接口当前主要用于统一前后端退出动作，接口返回成功后前端需要主动删除本地 token。",
            "2. 由于系统采用 JWT 无状态认证，当前版本不会在服务端维护会话黑名单。",
            "3. 若未来接入强制下线、单点登录或刷新令牌机制，可继续沿用该接口做扩展。",
        ],
    ),
    endpoint(
        "获取当前登录用户信息接口",
        "鉴权查询接口",
        "获取当前登录用户信息",
        "/api/auth/me",
        "GET",
        "application/json",
        "需携带普通用户 Token 或管理员 Token",
        "根据当前令牌解析登录用户身份，并返回用户档案信息与实名状态。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"登录令牌，格式 {USER_TOKEN}"),
        ],
        ["GET /api/auth/me", f"Authorization: {USER_TOKEN}"],
        USER_PROFILE_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "userId": 1001,
                    "phone": "13800138000",
                    "nickname": "小陈",
                    "avatarUrl": "http://localhost:9001/api/files/view/avatar.png",
                    "role": "USER",
                    "realNameStatus": "APPROVED",
                    "gender": "男",
                    "contactPhone": "13800138000",
                    "bio": "校园跑腿达人",
                },
            }
        ),
        [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
        [
            "1. 该接口是前端页面初始化时校验登录态的核心查询接口，适合在应用启动或页面刷新时调用。",
            "2. 后端通过 JWT 解析当前用户 ID，再从用户表、资料表与实名表拼装统一的用户视图对象。",
            "3. 前端可利用 role、realNameStatus 等字段控制菜单权限、资料页显示和实名认证引导。",
        ],
    ),
    endpoint(
        "获取个人资料接口",
        "用户资料接口",
        "获取个人资料",
        "/api/user/profile",
        "GET",
        "application/json",
        "需携带普通用户 Token",
        "返回当前普通用户个人资料，与 /auth/me 返回结构保持一致。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
        ],
        ["GET /api/user/profile", f"Authorization: {USER_TOKEN}"],
        USER_PROFILE_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "userId": 1001,
                    "phone": "13800138000",
                    "nickname": "小陈",
                    "avatarUrl": None,
                    "role": "USER",
                    "realNameStatus": "PENDING",
                    "gender": "女",
                    "contactPhone": "13800138000",
                    "bio": "喜欢结伴学习",
                },
            }
        ),
        [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
        [
            "1. 该接口面向普通用户资料页使用，路径语义更贴近“用户中心”。",
            "2. 虽然返回结构与 /auth/me 一致，但业务上更适合在用户资料模块中调用。",
            "3. 当前系统通过统一 VO 保证不同页面读取到的资料结构完全一致，减少前端分支判断。",
        ],
    ),
    endpoint(
        "更新个人资料接口",
        "用户资料维护接口",
        "更新个人资料",
        "/api/user/profile",
        "PUT",
        "application/json",
        "需携带普通用户 Token",
        "修改昵称、头像、性别、联系方式和个人简介，并返回最新资料。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            param_row("Body", "nickname", "string", "是", "昵称，最长 64 字符"),
            param_row("Body", "avatarUrl", "string", "否", "头像访问地址"),
            param_row("Body", "gender", "string", "否", "性别，最长 16 字符"),
            param_row("Body", "contactPhone", "string", "否", "联系手机号，最长 64 字符"),
            param_row("Body", "bio", "string", "否", "个人简介，最长 255 字符"),
        ],
        success_lines(
            {
                "nickname": "小陈",
                "avatarUrl": "http://localhost:9001/api/files/view/avatar.png",
                "gender": "女",
                "contactPhone": "13800138000",
                "bio": "可以帮取快递，也想找学习搭子",
            }
        ),
        USER_PROFILE_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "userId": 1001,
                    "phone": "13800138000",
                    "nickname": "小陈",
                    "avatarUrl": "http://localhost:9001/api/files/view/avatar.png",
                    "role": "USER",
                    "realNameStatus": "APPROVED",
                    "gender": "女",
                    "contactPhone": "13800138000",
                    "bio": "可以帮取快递，也想找学习搭子",
                },
            }
        ),
        [("失败返回（参数不合法）", fail_lines(400, "请求参数不合法"))],
        [
            "1. 该接口使用 PUT 表示对当前登录用户资料做整体更新，符合资源更新语义。",
            "2. 后端会同步更新 user 表中的 nickname、avatarUrl，以及 user_profile 表中的扩展资料。",
            "3. 更新完成后直接返回最新资料对象，前端可据此实时刷新个人中心与顶部头像信息。",
        ],
    ),
    endpoint(
        "提交实名认证接口",
        "实名认证接口",
        "提交实名认证",
        "/api/verification/submit",
        "POST",
        "application/json",
        "需携带普通用户 Token",
        "提交或重新提交实名信息与身份证图片，进入待审核状态。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            param_row("Body", "realName", "string", "是", "真实姓名"),
            param_row("Body", "idCardNo", "string", "是", "身份证号码"),
            param_row("Body", "frontFileId", "number", "是", "身份证正面图片文件 ID"),
            param_row("Body", "backFileId", "number", "是", "身份证反面图片文件 ID"),
        ],
        success_lines(
            {
                "realName": "陈一一",
                "idCardNo": "440101199901018888",
                "frontFileId": 301,
                "backFileId": 302,
            }
        ),
        VERIFICATION_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "id": 12,
                    "realName": "陈一一",
                    "idCardNo": "440101199901018888",
                    "frontFileId": 301,
                    "frontFileUrl": "http://localhost:9001/api/files/view/id-front.png",
                    "backFileId": 302,
                    "backFileUrl": "http://localhost:9001/api/files/view/id-back.png",
                    "status": "PENDING",
                    "rejectReason": None,
                    "reviewedBy": None,
                    "reviewedAt": None,
                },
            }
        ),
        [("失败返回（参数不合法）", fail_lines(400, "请求参数不合法"))],
        [
            "1. 用户首次提交和驳回后再次提交都复用该接口，后端会按 userId 查找并更新同一条实名记录。",
            "2. 每次提交都会把状态重置为 PENDING，同时清空历史驳回原因与审核信息。",
            "3. 建议前端先完成图片上传，再把返回的 fileId 填入实名提交请求体中。",
        ],
    ),
    endpoint(
        "获取当前实名认证信息接口",
        "实名认证查询接口",
        "获取当前实名认证信息",
        "/api/verification/current",
        "GET",
        "application/json",
        "需携带普通用户 Token",
        "获取当前登录用户的实名提交记录与审核状态；若从未提交则返回 data=null。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
        ],
        ["GET /api/verification/current", f"Authorization: {USER_TOKEN}"],
        VERIFICATION_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "id": 12,
                    "realName": "陈一一",
                    "idCardNo": "440101199901018888",
                    "frontFileId": 301,
                    "frontFileUrl": "http://localhost:9001/api/files/view/id-front.png",
                    "backFileId": 302,
                    "backFileUrl": "http://localhost:9001/api/files/view/id-back.png",
                    "status": "APPROVED",
                    "rejectReason": None,
                    "reviewedBy": 9001,
                    "reviewedAt": "2026-06-04T10:21:00",
                },
            }
        ),
        [("成功返回（未提交过实名认证）", success_lines({"code": 0, "message": "success", "data": None}))],
        [
            "1. 该接口用于实名页面回显当前状态，避免用户重复提交或重复上传证件图片。",
            "2. 当前实现对“没有实名记录”的场景返回 code=0 且 data=null，而不是 404。",
            "3. 前端可结合 status 和 rejectReason 展示“待审核”“已通过”“驳回原因”等差异化 UI。",
        ],
    ),
    endpoint(
        "获取跑腿规则接口",
        "跑腿基础配置接口",
        "获取跑腿规则",
        "/api/errand/rules",
        "GET",
        "application/json",
        "需携带普通用户 Token",
        "获取当前跑腿费用规则，包括加急费和易碎费。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
        ],
        ["GET /api/errand/rules", f"Authorization: {USER_TOKEN}"],
        ERRAND_RULE_FIELDS,
        success_lines({"code": 0, "message": "success", "data": {"urgentFee": 2.0, "fragileFee": 3.0}}),
        [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
        [
            "1. 跑腿下单页在计算预估费用前应先调用该接口获取后台配置，而不是写死在前端。",
            "2. 当前规则保存在应用配置中，因此接口返回速度快且结构简单。",
            "3. 管理后台修改规则后，移动端再次请求该接口即可拿到新的计费参数。",
        ],
    ),
    endpoint(
        "获取跑腿广场列表接口",
        "跑腿列表接口",
        "获取跑腿广场列表",
        "/api/errand/orders/square",
        "GET",
        "application/json",
        "需携带普通用户 Token",
        "分页获取仍在公开展示且未超时的跑腿订单列表，支持关键字搜索。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            param_row("Query", "keyword", "string", "否", "关键字，可匹配服务类型、取货地址、送达地址、备注"),
            param_row("Query", "current", "number", "否", "页码，默认 1"),
            param_row("Query", "size", "number", "否", "每页条数，默认 10"),
        ],
        [
            "GET /api/errand/orders/square?keyword=快递&current=1&size=10",
            f"Authorization: {USER_TOKEN}",
        ],
        ERRAND_ORDER_PAGE_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "current": 1,
                    "size": 10,
                    "total": 1,
                    "pages": 1,
                    "records": [
                        {
                            "id": 88,
                            "orderNo": "ET20260604120100123",
                            "serviceType": "PICKUP",
                            "pickupAddress": "菜鸟驿站",
                            "deliveryAddress": "2 号宿舍楼 305",
                            "totalFee": 8.0,
                            "status": "PUBLISHED",
                            "publisher": {"nickname": "小陈"},
                            "runner": None,
                            "conversationId": None,
                            "createdAt": "2026-06-04T12:01:00",
                        }
                    ],
                },
            }
        ),
        [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
        [
            "1. 查询前服务端会自动把已过接单截止时间的公开订单转为 EXPIRED，保证广场数据实时有效。",
            "2. 该接口返回的是分页对象，前端需要从 data.records 中读取列表，从 data.total 中读取总数。",
            "3. canAccept 等权限布尔值会基于当前登录用户身份计算，方便前端直接控制按钮状态。",
        ],
    ),
    endpoint(
        "创建跑腿订单接口",
        "跑腿发布接口",
        "创建跑腿订单",
        "/api/errand/orders",
        "POST",
        "application/json",
        "需携带普通用户 Token",
        "发布新的跑腿订单，系统会自动生成订单号、接单截止时间与总费用。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            param_row("Body", "serviceType", "string", "是", "服务类型：PICKUP/DELIVERY/PURCHASE/PRINT"),
            param_row("Body", "pickupAddress", "string", "是", "取货地址，最长 255 字符"),
            param_row("Body", "deliveryAddress", "string", "是", "送达地址，最长 255 字符"),
            param_row("Body", "pickupTimeText", "string", "否", "期望取货时间说明，最长 100 字符"),
            param_row("Body", "detailContent", "string", "是", "详细需求，最长 500 字符"),
            param_row("Body", "remark", "string", "否", "补充备注，最长 500 字符"),
            param_row("Body", "baseFee", "number", "是", "基础费用，最小值 0.01"),
            param_row("Body", "urgentFlag", "boolean", "是", "是否加急"),
            param_row("Body", "fragileFlag", "boolean", "是", "是否易碎"),
            param_row("Body", "attachmentFileIds", "number[]", "是", "订单附件文件 ID 列表，可为空数组"),
        ],
        success_lines(
            {
                "serviceType": "PICKUP",
                "pickupAddress": "菜鸟驿站",
                "deliveryAddress": "2 号宿舍楼 305",
                "pickupTimeText": "今晚 8 点前",
                "detailContent": "帮忙取快递，编号 A1234",
                "remark": "到楼下电话联系",
                "baseFee": 5.0,
                "urgentFlag": True,
                "fragileFlag": False,
                "attachmentFileIds": [410, 411],
            }
        ),
        ERRAND_ORDER_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "id": 88,
                    "orderNo": "ET20260604120100123",
                    "serviceType": "PICKUP",
                    "pickupAddress": "菜鸟驿站",
                    "deliveryAddress": "2 号宿舍楼 305",
                    "pickupTimeText": "今晚 8 点前",
                    "detailContent": "帮忙取快递，编号 A1234",
                    "remark": "到楼下电话联系",
                    "baseFee": 5.0,
                    "urgentFee": 2.0,
                    "fragileFee": 0.0,
                    "totalFee": 7.0,
                    "urgentFlag": True,
                    "fragileFlag": False,
                    "status": "PUBLISHED",
                    "acceptDeadline": "2026-06-04T16:01:00",
                    "acceptedAt": None,
                    "pickedUpAt": None,
                    "deliveredAt": None,
                    "completedAt": None,
                    "cancelledAt": None,
                    "cancelReason": None,
                    "publicVisible": True,
                    "publisher": {"userId": 1001, "nickname": "小陈", "phone": "13800138000"},
                    "runner": None,
                    "attachments": [{"fileId": 410, "url": "http://localhost:9001/api/files/view/a.png", "originalName": "a.png"}],
                    "relatedReports": [],
                    "conversationId": None,
                    "canAccept": False,
                    "canCancel": True,
                    "canStart": False,
                    "canDeliver": False,
                    "canComplete": False,
                    "canReport": True,
                    "createdAt": "2026-06-04T12:01:00",
                },
            }
        ),
        [
            ("失败返回（服务类型不合法）", fail_lines(400, "服务类型不合法")),
            ("失败返回（参数不合法）", fail_lines(400, "请求参数不合法")),
        ],
        [
            "1. 下单前前端应先上传图片附件并获取 fileId，再连同业务字段一起提交。",
            "2. 后端会根据 urgentFlag、fragileFlag 自动计算附加费和 totalFee，并把订单状态初始化为 PUBLISHED。",
            "3. 新订单默认公开展示 4 小时，接单截止时间由服务端统一生成，避免被前端篡改。",
        ],
    ),
    endpoint(
        "获取跑腿订单详情接口",
        "跑腿详情接口",
        "获取跑腿订单详情",
        "/api/errand/orders/{id}",
        "GET",
        "application/json",
        "需携带普通用户 Token",
        "查看订单详情；公开订单可查看，非公开订单仅参与方可查看。",
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            param_row("Path", "id", "number", "是", "跑腿订单 ID"),
        ],
        ["GET /api/errand/orders/88", f"Authorization: {USER_TOKEN}"],
        ERRAND_ORDER_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "id": 88,
                    "orderNo": "ET20260604120100123",
                    "serviceType": "PICKUP",
                    "pickupAddress": "菜鸟驿站",
                    "deliveryAddress": "2 号宿舍楼 305",
                    "pickupTimeText": "今晚 8 点前",
                    "detailContent": "帮忙取快递，编号 A1234",
                    "remark": "到楼下电话联系",
                    "baseFee": 5.0,
                    "urgentFee": 2.0,
                    "fragileFee": 0.0,
                    "totalFee": 7.0,
                    "urgentFlag": True,
                    "fragileFlag": False,
                    "status": "PUBLISHED",
                    "acceptDeadline": "2026-06-04T16:01:00",
                    "publicVisible": True,
                    "publisher": {"userId": 1001, "nickname": "小陈", "phone": "13800138000"},
                    "runner": None,
                    "attachments": [],
                    "relatedReports": [],
                    "conversationId": None,
                    "canAccept": True,
                    "canCancel": False,
                    "canStart": False,
                    "canDeliver": False,
                    "canComplete": False,
                    "canReport": False,
                    "createdAt": "2026-06-04T12:01:00",
                },
            }
        ),
        [
            ("失败返回（订单不存在）", fail_lines(400, "跑腿订单不存在")),
            ("失败返回（无权查看）", fail_lines(403, "无权查看该订单")),
        ],
        [
            "1. 服务端在返回详情前也会先做过期处理，保证超时未接单的订单不会继续显示为 PUBLISHED。",
            "2. 如果订单仍处于公开状态，任何登录用户都可查看；否则只有发布人和接单人可访问。",
            "3. 返回结果包含大量前端直接可用的按钮权限位，例如 canAccept、canCancel、canComplete。",
        ],
    ),
]


def add_errand_action_endpoint(order: int, title: str, path: str, action_name: str, action_desc: str, success_status: str, failure_message: str, logic_lines: list[str]) -> dict:
    return endpoint(
        title,
        "跑腿状态流转接口",
        action_name,
        path,
        "POST",
        "application/json",
        "需携带普通用户 Token",
        action_desc,
        [
            ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
            param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            param_row("Path", "id", "number", "是", "跑腿订单 ID"),
        ],
        [f"POST {path.replace('{id}', '88')}", f"Authorization: {USER_TOKEN}", "", "{}"],
        ERRAND_ORDER_FIELDS,
        success_lines(
            {
                "code": 0,
                "message": "success",
                "data": {
                    "id": 88,
                    "orderNo": "ET20260604120100123",
                    "serviceType": "PICKUP",
                    "pickupAddress": "菜鸟驿站",
                    "deliveryAddress": "2 号宿舍楼 305",
                    "totalFee": 7.0,
                    "status": success_status,
                    "publisher": {"userId": 1001, "nickname": "小陈", "phone": "13800138000"},
                    "runner": {"userId": 1002, "nickname": "小林", "phone": "13900001111"},
                    "conversationId": 501,
                    "canAccept": False,
                    "canCancel": False,
                    "canStart": False,
                    "canDeliver": False,
                    "canComplete": success_status == "DELIVERING",
                    "canReport": True,
                    "createdAt": "2026-06-04T12:01:00",
                },
            }
        ),
        [
            ("失败返回（订单不存在）", fail_lines(400, "跑腿订单不存在")),
            ("失败返回（业务状态不允许）", fail_lines(400, failure_message)),
        ],
        logic_lines,
    )


endpoints.extend(
    [
        add_errand_action_endpoint(
            0,
            "承接跑腿订单接口",
            "/api/errand/orders/{id}/accept",
            "承接跑腿订单",
            "由其他用户承接公开跑腿订单，并创建双方沟通会话。",
            "ACCEPTED",
            "订单已被承接",
            [
                "1. 接单人不能承接自己发布的订单，且订单必须仍为公开发布状态。",
                "2. 接单成功后系统会自动创建或补齐订单私聊会话，并推送“订单已被承接”通知消息。",
                "3. 订单被承接后会从广场下架，runner 信息、conversationId 等字段开始生效。",
            ],
        ),
        add_errand_action_endpoint(
            0,
            "开始执行跑腿订单接口",
            "/api/errand/orders/{id}/start",
            "开始执行跑腿订单",
            "由接单人将订单状态从 ACCEPTED 推进到 IN_PROGRESS。",
            "IN_PROGRESS",
            "当前订单不可开始履约",
            [
                "1. 只有当前订单的接单人可以调用该接口，发布人无权代替执行。",
                "2. 接口成功后状态改为 IN_PROGRESS，并向订单会话推送进度通知。",
                "3. 前端通常在双方已经确认需求后再触发此动作，以区分“已接单”和“已开始执行”。",
            ],
        ),
        add_errand_action_endpoint(
            0,
            "提交送达跑腿订单接口",
            "/api/errand/orders/{id}/deliver",
            "提交送达跑腿订单",
            "由接单人标记订单进入派送/处理中状态。",
            "DELIVERING",
            "当前订单不可更新为处理中",
            [
                "1. 只有接单人可提交送达，且订单当前必须处于 ACCEPTED 或 IN_PROGRESS。",
                "2. 后端会同时写入 pickedUpAt 与 deliveredAt 时间，表示已进入实际履约阶段。",
                "3. 前端在该状态后应引导发布人确认收货或确认完成。",
            ],
        ),
        add_errand_action_endpoint(
            0,
            "确认完成跑腿订单接口",
            "/api/errand/orders/{id}/complete",
            "确认完成跑腿订单",
            "由发布人确认订单已完成并结束本次履约。",
            "COMPLETED",
            "当前订单不可确认完成",
            [
                "1. 该接口只能由发布人调用，防止接单人自行结束订单。",
                "2. 订单必须已经进入 DELIVERING，后端才允许推进到 COMPLETED。",
                "3. 完成后系统会向会话中推送完成通知，后续如出现争议可通过举报流程介入。",
            ],
        ),
        add_errand_action_endpoint(
            0,
            "取消跑腿订单接口",
            "/api/errand/orders/{id}/cancel",
            "取消跑腿订单",
            "由发布人主动取消尚未被承接的订单。",
            "CANCELLED",
            "当前订单不可取消",
            [
                "1. 只有发布人可以取消订单，且订单必须仍处于 PUBLISHED 未承接状态。",
                "2. 取消后订单会从广场下架，cancelReason 固定写入“发布人主动取消”。",
                "3. 对于已承接或已履约中的订单，需走管理员介入或举报处理流程，而不是调用该接口。",
            ],
        ),
        endpoint(
            "获取我发布的跑腿订单列表接口",
            "跑腿个人列表接口",
            "获取我发布的跑腿订单列表",
            "/api/errand/orders/my/published",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "按创建时间倒序返回当前用户自己发布的全部跑腿订单。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            ],
            ["GET /api/errand/orders/my/published", f"Authorization: {USER_TOKEN}"],
            ERRAND_ORDER_LIST_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": [
                        {
                            "id": 88,
                            "orderNo": "ET20260604120100123",
                            "serviceType": "PICKUP",
                            "pickupAddress": "菜鸟驿站",
                            "deliveryAddress": "2 号宿舍楼 305",
                            "totalFee": 7.0,
                            "status": "PUBLISHED",
                            "publisher": {"nickname": "小陈"},
                            "runner": None,
                            "conversationId": None,
                            "createdAt": "2026-06-04T12:01:00",
                        }
                    ],
                }
            ),
            [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 该接口适合“我发布的”列表页使用，不需要额外分页参数。",
                "2. 服务端会在查询前先处理超时订单，因此列表中的状态始终是最新状态。",
                "3. 返回的每条订单仍带有 conversationId、权限位等字段，便于列表直接进入沟通或详情页。",
            ],
        ),
        endpoint(
            "获取我承接的跑腿订单列表接口",
            "跑腿个人列表接口",
            "获取我承接的跑腿订单列表",
            "/api/errand/orders/my/accepted",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "按创建时间倒序返回当前用户承接过的跑腿订单。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            ],
            ["GET /api/errand/orders/my/accepted", f"Authorization: {USER_TOKEN}"],
            ERRAND_ORDER_LIST_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": [
                        {
                            "id": 66,
                            "orderNo": "ET20260603110100456",
                            "serviceType": "DELIVERY",
                            "pickupAddress": "打印店",
                            "deliveryAddress": "图书馆门口",
                            "totalFee": 6.0,
                            "status": "IN_PROGRESS",
                            "publisher": {"nickname": "小王"},
                            "runner": {"nickname": "小林"},
                            "conversationId": 501,
                            "createdAt": "2026-06-03T11:01:00",
                        }
                    ],
                }
            ),
            [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 该接口面向接单人视角，通常用于“我承接的”或“我的任务”页面。",
                "2. 列表中可直接读取当前状态，决定是展示“开始执行”“提交送达”还是“查看会话”等动作。",
                "3. 由于返回的是当前用户参与的订单，非公开订单信息也会完整返回。",
            ],
        ),
        endpoint(
            "获取跑腿订单会话信息接口",
            "跑腿沟通接口",
            "获取跑腿订单会话信息",
            "/api/errand/orders/{id}/conversation",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "获取或确保当前跑腿订单对应的私聊会话 ID，仅发布人和接单人可调用。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("Path", "id", "number", "是", "跑腿订单 ID"),
            ],
            ["GET /api/errand/orders/88/conversation", f"Authorization: {USER_TOKEN}"],
            ERRAND_CONVERSATION_FIELDS,
            success_lines({"code": 0, "message": "success", "data": {"conversationId": 501}}),
            [("失败返回（无权访问）", fail_lines(403, "无权访问该订单会话"))],
            [
                "1. 该接口用于详情页跳转聊天前获取 conversationId，前端无需自己推断会话编号。",
                "2. 如果订单已有会话，接口直接返回原会话；如果已满足会话创建条件，系统会自动补齐成员关系。",
                "3. 非参与方即使知道订单 ID 也不能通过该接口获得会话信息，服务端会返回 403。",
            ],
        ),
        endpoint(
            "创建举报接口",
            "举报提交接口",
            "创建举报",
            "/api/reports",
            "POST",
            "application/json",
            "需携带普通用户 Token",
            "提交举报工单，支持附带附件文件，并对频率和重复举报做限制。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("Body", "module", "string", "是", "业务模块，例如 errand"),
                param_row("Body", "targetType", "string", "是", "目标类型，例如 order"),
                param_row("Body", "targetId", "number", "否", "被举报目标 ID"),
                param_row("Body", "reportType", "string", "是", "举报类型，最长 50 字符"),
                param_row("Body", "description", "string", "是", "举报说明，最长 500 字符"),
                param_row("Body", "contactPhone", "string", "否", "联系电话，最长 64 字符"),
                param_row("Body", "attachmentFileIds", "number[]", "否", "附件文件 ID 列表"),
            ],
            success_lines(
                {
                    "module": "errand",
                    "targetType": "order",
                    "targetId": 88,
                    "reportType": "未按约履约",
                    "description": "接单后长时间不回复，且未按约送达。",
                    "contactPhone": "13800138000",
                    "attachmentFileIds": [601],
                }
            ),
            REPORT_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 900,
                        "module": "errand",
                        "targetType": "order",
                        "targetId": 88,
                        "reportType": "未按约履约",
                        "description": "接单后长时间不回复，且未按约送达。",
                        "contactPhone": "13800138000",
                        "status": "PENDING",
                        "handleRemark": None,
                        "attachments": [{"fileId": 601, "url": "http://localhost:9001/api/files/view/proof.png", "originalName": "proof.png"}],
                        "handledBy": None,
                        "handledAt": None,
                        "createdAt": "2026-06-04T13:20:00",
                    },
                }
            ),
            [
                ("失败返回（重复举报）", fail_lines(400, "该内容已有处理中举报，请勿重复提交")),
                ("失败返回（提交过频）", fail_lines(400, "举报提交过于频繁，请 10 分钟后再试")),
            ],
            [
                "1. 系统会限制同一目标短时间内重复举报，也会限制用户 10 分钟内最多提交 3 次举报。",
                "2. 如果用户最近 30 天被驳回的举报过多，接口会临时限制其继续提交举报。",
                "3. 当举报目标是跑腿订单且管理员后续将其状态处理为 PROCESSING 时，订单会自动进入争议状态。",
            ],
        ),
        endpoint(
            "获取我的举报列表接口",
            "举报查询接口",
            "获取我的举报列表",
            "/api/reports",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "按创建时间倒序返回当前用户提交的全部举报工单。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            ],
            ["GET /api/reports", f"Authorization: {USER_TOKEN}"],
            [
                field_row("data[].id", "number", "举报单 ID"),
                field_row("data[].module", "string", "业务模块"),
                field_row("data[].targetType", "string", "目标类型"),
                field_row("data[].targetId", "number", "目标 ID"),
                field_row("data[].reportType", "string", "举报类型"),
                field_row("data[].description", "string", "举报说明"),
                field_row("data[].contactPhone", "string", "联系电话"),
                field_row("data[].status", "string", "处理状态"),
                field_row("data[].handleRemark", "string", "处理备注"),
                field_row("data[].attachments[].fileId", "number", "附件文件 ID"),
                field_row("data[].handledBy", "number", "处理管理员 ID"),
                field_row("data[].handledAt", "string(datetime)", "处理时间"),
                field_row("data[].createdAt", "string(datetime)", "创建时间"),
            ],
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": [
                        {
                            "id": 900,
                            "module": "errand",
                            "targetType": "order",
                            "targetId": 88,
                            "reportType": "未按约履约",
                            "description": "接单后长时间不回复，且未按约送达。",
                            "contactPhone": "13800138000",
                            "status": "PENDING",
                            "handleRemark": None,
                            "attachments": [],
                            "handledBy": None,
                            "handledAt": None,
                            "createdAt": "2026-06-04T13:20:00",
                        }
                    ],
                }
            ),
            [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 该接口返回当前登录用户自己的举报单，不会暴露其他用户的举报数据。",
                "2. 前端可以根据 status 展示“待处理”“处理中”“已解决”“已驳回”等状态标签。",
                "3. 附件和处理信息都会跟随列表返回，便于在列表页或卡片视图中直接展示关键信息。",
            ],
        ),
        endpoint(
            "获取举报详情接口",
            "举报查询接口",
            "获取举报详情",
            "/api/reports/{id}",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "获取当前用户自己提交的举报详情；若不是本人举报则返回 data=null。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("Path", "id", "number", "是", "举报单 ID"),
            ],
            ["GET /api/reports/900", f"Authorization: {USER_TOKEN}"],
            REPORT_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 900,
                        "module": "errand",
                        "targetType": "order",
                        "targetId": 88,
                        "reportType": "未按约履约",
                        "description": "接单后长时间不回复，且未按约送达。",
                        "contactPhone": "13800138000",
                        "status": "PROCESSING",
                        "handleRemark": "已联系双方核实",
                        "attachments": [],
                        "handledBy": 9001,
                        "handledAt": "2026-06-04T15:10:00",
                        "createdAt": "2026-06-04T13:20:00",
                    },
                }
            ),
            [("成功返回（非本人举报或记录不存在）", success_lines({"code": 0, "message": "success", "data": None}))],
            [
                "1. 该接口会校验举报单 submitterId 是否与当前用户一致，不一致时不会抛错，而是返回 data=null。",
                "2. 这种设计可以降低信息泄露风险，同时简化前端对“详情不存在”和“无权限”的统一处理。",
                "3. 若需要强提示，前端可在收到 data=null 后跳回列表页并给出友好提示。",
            ],
        ),
        endpoint(
            "获取会话列表接口",
            "消息会话接口",
            "获取会话列表",
            "/api/conversations",
            "GET",
            "application/json",
            "需携带普通用户 Token 或管理员 Token",
            "获取当前登录用户所属的全部会话列表，并带上未读数和最后一条消息。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", "登录令牌，普通用户或管理员均可"),
            ],
            ["GET /api/conversations", f"Authorization: {USER_TOKEN}"],
            CONVERSATION_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": [
                        {
                            "id": 501,
                            "type": "PRIVATE",
                            "title": "跑腿订单 ET20260604120100123",
                            "unreadCount": 2,
                            "lastMessage": {
                                "id": 1200,
                                "type": "TEXT",
                                "content": "我已经到驿站了",
                                "createdAt": "2026-06-04T14:01:00",
                            },
                            "updatedAt": "2026-06-04T14:01:00",
                        }
                    ],
                }
            ),
            [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 会话列表由会话成员关系表驱动，因此只会返回当前用户参与的会话。",
                "2. 返回结果已经按更新时间倒序排序，前端通常可直接用于消息首页渲染。",
                "3. unreadCount 和 lastMessage 是会话列表页最关键的摘要字段，可直接驱动红点和预览文案。",
            ],
        ),
        endpoint(
            "获取会话消息列表接口",
            "消息明细接口",
            "获取会话消息列表",
            "/api/conversations/{id}/messages",
            "GET",
            "application/json",
            "需携带普通用户 Token 或管理员 Token",
            "获取指定会话的全部消息，按消息 ID 升序返回。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", "登录令牌，普通用户或管理员均可"),
                param_row("Path", "id", "number", "是", "会话 ID"),
            ],
            ["GET /api/conversations/501/messages", f"Authorization: {USER_TOKEN}"],
            MESSAGES_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": [
                        {
                            "id": 1198,
                            "conversationId": 501,
                            "senderId": 1001,
                            "type": "TEXT",
                            "content": "你好，现在可以帮我取件吗？",
                            "createdAt": "2026-06-04T13:50:00",
                        },
                        {
                            "id": 1200,
                            "conversationId": 501,
                            "senderId": 1002,
                            "type": "TEXT",
                            "content": "我已经到驿站了",
                            "createdAt": "2026-06-04T14:01:00",
                        },
                    ],
                }
            ),
            [
                ("失败返回（会话不存在或无权限）", fail_lines(400, "无权访问该会话")),
            ],
            [
                "1. 只有会话成员才能拉取消息内容，服务端会基于 conversation_member 表做权限校验。",
                "2. 当前接口返回完整历史消息列表，适合消息量不大的校园场景；后续可扩展分页。",
                "3. 消息按 ID 升序返回，前端无需再次排序，可直接按时间轴从上到下渲染。",
            ],
        ),
        endpoint(
            "发送消息接口",
            "消息发送接口",
            "发送消息",
            "/api/messages/send",
            "POST",
            "application/json",
            "需携带普通用户 Token 或管理员 Token",
            "向指定会话发送文本或图片消息，并更新其他成员未读数。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", "登录令牌，普通用户或管理员均可"),
                param_row("Body", "conversationId", "number", "是", "目标会话 ID"),
                param_row("Body", "type", "string", "是", "消息类型，TEXT 或 IMAGE"),
                param_row("Body", "content", "string", "是", "消息内容；图片消息通常放图片 URL"),
            ],
            success_lines({"conversationId": 501, "type": "TEXT", "content": "我已经到驿站了"}),
            MESSAGE_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 1200,
                        "conversationId": 501,
                        "senderId": 1002,
                        "type": "TEXT",
                        "content": "我已经到驿站了",
                        "createdAt": "2026-06-04T14:01:00",
                    },
                }
            ),
            [
                ("失败返回（无权访问该会话）", fail_lines(400, "无权访问该会话")),
                ("失败返回（会话不存在）", fail_lines(400, "会话不存在")),
            ],
            [
                "1. 服务端会先校验发送人是否具备该会话的发言权限，再写入消息表。",
                "2. 写入成功后系统会更新会话 lastMessageId，并给其他成员增加未读数、推送 WebSocket 新消息事件。",
                "3. 管理员在 SYSTEM 会话中即使不是 conversation_member，也可向系统通知会话发送消息。",
            ],
        ),
        endpoint(
            "标记会话已读接口",
            "消息状态接口",
            "标记会话已读",
            "/api/conversations/{id}/read",
            "POST",
            "application/json",
            "需携带普通用户 Token 或管理员 Token",
            "将指定会话的未读数清零，并记录最后已读消息。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", "登录令牌，普通用户或管理员均可"),
                param_row("Path", "id", "number", "是", "会话 ID"),
            ],
            ["POST /api/conversations/501/read", f"Authorization: {USER_TOKEN}", "", "{}"],
            [],
            success_lines({"code": 0, "message": "success", "data": None}),
            [("失败返回（无权访问该会话）", fail_lines(400, "无权访问该会话"))],
            [
                "1. 用户进入聊天详情并拉取消息后，通常立即调用该接口清除红点。",
                "2. 后端会把 unreadCount 置 0，并把 lastReadMessageId 更新为当前会话最后一条消息。",
                "3. 同时系统会向当前用户推送 conversation.unread WebSocket 事件，用于实时同步角标状态。",
            ],
        ),
        endpoint(
            "获取未读消息统计接口",
            "消息统计接口",
            "获取未读消息统计",
            "/api/messages/unread",
            "GET",
            "application/json",
            "需携带普通用户 Token 或管理员 Token",
            "返回当前用户全部会话的未读消息总数。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", "登录令牌，普通用户或管理员均可"),
            ],
            ["GET /api/messages/unread", f"Authorization: {USER_TOKEN}"],
            UNREAD_FIELDS,
            success_lines({"code": 0, "message": "success", "data": {"totalUnread": 5}}),
            [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 该接口适合首页角标、Tab 红点和全局消息提醒使用。",
                "2. 服务端会遍历当前用户的会话成员记录并累加 unreadCount。",
                "3. 如果前端已接入 WebSocket，也可将此接口作为首次冷启动时的初始化统计来源。",
            ],
        ),
        endpoint(
            "上传图片接口",
            "文件上传接口",
            "上传图片",
            "/api/files/upload/image",
            "POST",
            "multipart/form-data",
            "需携带普通用户 Token 或管理员 Token",
            "上传图片文件并返回 fileId 和公开访问地址，供头像、举报、跑腿附件等场景复用。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", "登录令牌，普通用户或管理员均可"),
                param_row("FormData", "file", "file", "是", "待上传的图片文件"),
                param_row("FormData", "bizType", "string", "否", "业务类型，默认 common"),
            ],
            [
                "POST /api/files/upload/image",
                f"Authorization: {USER_TOKEN}",
                "Content-Type: multipart/form-data",
                "",
                "file = avatar.png",
                "bizType = avatar",
            ],
            UPLOAD_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "fileId": 410,
                        "url": "http://localhost:9001/api/files/view/2e3b3e7d.png",
                        "originalName": "avatar.png",
                    },
                }
            ),
            [
                ("失败返回（未选择文件）", fail_lines(400, "请选择要上传的文件")),
                ("失败返回（上传失败）", fail_lines(500, "文件上传失败")),
            ],
            [
                "1. 当前接口只负责把文件落盘并登记数据库元数据，不直接绑定具体业务对象。",
                "2. 业务侧后续通过 fileId 与举报、跑腿订单、实名资料等表建立关联关系。",
                "3. 返回的 url 可直接用于前端预览，而 fileId 应继续保存以便提交业务接口。",
            ],
        ),
        endpoint(
            "美妆商品实拍图上传接口",
            "美妆模块复用接口",
            "上传美妆商品实拍图",
            "/api/files/upload/image",
            "POST",
            "multipart/form-data",
            "需携带普通用户 Token",
            "在美妆发布页上传商品实拍图，通常配合 bizType=beauty_product 使用。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("FormData", "file", "file", "是", "商品实拍图文件"),
                param_row("FormData", "bizType", "string", "否", "业务类型，建议传 beauty_product"),
            ],
            [
                "POST /api/files/upload/image",
                f"Authorization: {USER_TOKEN}",
                "Content-Type: multipart/form-data",
                "",
                "file = beauty-good.jpg",
                "bizType = beauty_product",
            ],
            UPLOAD_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "fileId": 701,
                        "url": "http://localhost:9001/api/files/view/beauty-good.jpg",
                        "originalName": "beauty-good.jpg",
                    },
                }
            ),
            [
                ("失败返回（未选择文件）", fail_lines(400, "请选择要上传的文件")),
                ("失败返回（上传失败）", fail_lines(500, "文件上传失败")),
            ],
            [
                "1. 美妆发布页的商品展示图需要先走该接口上传，再把返回的 url 与 fileId 用于页面预览和后续业务处理。",
                "2. 当前美妆发布能力仍以前端本地存储为主，因此 fileId 主要用于保留真实上传凭据，而不是立即写入商品表。",
                "3. 建议前端固定传 bizType=beauty_product，方便后续后端化后按业务类型做文件分类与清理。",
            ],
            section_header="美妆模块相关接口",
        ),
        endpoint(
            "美妆商品购买凭证上传接口",
            "美妆模块复用接口",
            "上传美妆商品购买凭证",
            "/api/files/upload/image",
            "POST",
            "multipart/form-data",
            "需携带普通用户 Token",
            "在美妆发布页上传订单购买凭证，通常配合 bizType=beauty_receipt 使用。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("FormData", "file", "file", "是", "订单购买凭证图片"),
                param_row("FormData", "bizType", "string", "否", "业务类型，建议传 beauty_receipt"),
            ],
            [
                "POST /api/files/upload/image",
                f"Authorization: {USER_TOKEN}",
                "Content-Type: multipart/form-data",
                "",
                "file = beauty-receipt.jpg",
                "bizType = beauty_receipt",
            ],
            UPLOAD_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "fileId": 702,
                        "url": "http://localhost:9001/api/files/view/beauty-receipt.jpg",
                        "originalName": "beauty-receipt.jpg",
                    },
                }
            ),
            [
                ("失败返回（未选择文件）", fail_lines(400, "请选择要上传的文件")),
                ("失败返回（上传失败）", fail_lines(500, "文件上传失败")),
            ],
            [
                "1. 购买凭证是当前美妆发布页的必传材料，用于证明商品来源，降低虚假种草或伪造商品信息的风险。",
                "2. 该接口与商品实拍图上传共用同一路径，只是 bizType 与业务语义不同。",
                "3. 当前版本里凭证上传成功后会先用于前端页面展示和本地存储，后续商品后端化后可直接复用该 fileId。",
            ],
        ),
        endpoint(
            "美妆商品举报接口",
            "美妆模块复用接口",
            "提交美妆商品举报",
            "/api/reports",
            "POST",
            "application/json",
            "需携带普通用户 Token",
            "对美妆商品发起举报，前端通常传 module=beauty、targetType=good、targetId=商品ID。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("Body", "module", "string", "是", "业务模块，固定传 beauty"),
                param_row("Body", "targetType", "string", "是", "目标类型，当前通常传 good"),
                param_row("Body", "targetId", "number", "是", "被举报的美妆商品 ID"),
                param_row("Body", "reportType", "string", "是", "举报类型，如 虚假宣传、盗图、价格异常"),
                param_row("Body", "description", "string", "是", "举报详细说明"),
                param_row("Body", "contactPhone", "string", "否", "联系人手机号"),
                param_row("Body", "attachmentFileIds", "number[]", "否", "补充证据附件文件 ID"),
            ],
            success_lines(
                {
                    "module": "beauty",
                    "targetType": "good",
                    "targetId": 5002,
                    "reportType": "虚假宣传",
                    "description": "商品描述与实拍严重不符，且疑似盗用他人图片。",
                    "contactPhone": "13800138000",
                    "attachmentFileIds": [703],
                }
            ),
            REPORT_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 980,
                        "module": "beauty",
                        "targetType": "good",
                        "targetId": 5002,
                        "reportType": "虚假宣传",
                        "description": "商品描述与实拍严重不符，且疑似盗用他人图片。",
                        "contactPhone": "13800138000",
                        "status": "PENDING",
                        "handleRemark": None,
                        "attachments": [{"fileId": 703, "url": "http://localhost:9001/api/files/view/beauty-proof.png", "originalName": "beauty-proof.png"}],
                        "handledBy": None,
                        "handledAt": None,
                        "createdAt": "2026-06-04T18:00:00",
                    },
                }
            ),
            [
                ("失败返回（重复举报）", fail_lines(400, "该内容已有处理中举报，请勿重复提交")),
                ("失败返回（提交过频）", fail_lines(400, "举报提交过于频繁，请 10 分钟后再试")),
            ],
            [
                "1. 美妆详情页和列表页的“举报”按钮当前都落到该正式后端接口，而不是前端本地逻辑。",
                "2. 虽然美妆商品本身还未完全后端化，但举报工单已经是真实入库、真实可被后台处理的业务数据。",
                "3. 若后续新增美妆商品后台表，现有举报参数结构仍可继续沿用，只需要保证 targetId 对应真实商品主键。",
            ],
        ),
        endpoint(
            "查询我的美妆举报列表接口",
            "美妆模块复用接口",
            "查询我的美妆举报列表",
            "/api/reports",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "查询当前用户提交的举报记录；前端需要自行筛选 module=beauty 的记录作为美妆举报列表。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            ],
            ["GET /api/reports", f"Authorization: {USER_TOKEN}"],
            [
                field_row("data[].id", "number", "举报单 ID"),
                field_row("data[].module", "string", "业务模块，筛选 beauty 即美妆举报"),
                field_row("data[].targetType", "string", "目标类型，当前常见为 good"),
                field_row("data[].targetId", "number", "被举报商品 ID"),
                field_row("data[].reportType", "string", "举报类型"),
                field_row("data[].description", "string", "举报说明"),
                field_row("data[].status", "string", "处理状态"),
                field_row("data[].attachments[].fileId", "number", "证据附件文件 ID"),
                field_row("data[].handledBy", "number", "处理管理员 ID"),
                field_row("data[].handledAt", "string(datetime)", "处理时间"),
                field_row("data[].createdAt", "string(datetime)", "创建时间"),
            ],
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": [
                        {
                            "id": 980,
                            "module": "beauty",
                            "targetType": "good",
                            "targetId": 5002,
                            "reportType": "虚假宣传",
                            "description": "商品描述与实拍严重不符，且疑似盗用他人图片。",
                            "contactPhone": "13800138000",
                            "status": "PENDING",
                            "handleRemark": None,
                            "attachments": [],
                            "handledBy": None,
                            "handledAt": None,
                            "createdAt": "2026-06-04T18:00:00",
                        }
                    ],
                }
            ),
            [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 当前接口返回的是“我的全部举报”，不是只返回美妆模块数据，因此前端要按 module=beauty 做一次筛选。",
                "2. 对于美妆模块来说，这个接口更适合作为“我的美妆举报记录”数据源复用。",
                "3. 如果后续需要单独的美妆举报列表接口，可以在后端增加按 module 过滤的查询参数。",
            ],
        ),
        endpoint(
            "查询美妆举报详情接口",
            "美妆模块复用接口",
            "查询美妆举报详情",
            "/api/reports/{id}",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "查看单条美妆举报详情，适用于用户从举报列表进入详情页。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("Path", "id", "number", "是", "举报单 ID"),
            ],
            ["GET /api/reports/980", f"Authorization: {USER_TOKEN}"],
            REPORT_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 980,
                        "module": "beauty",
                        "targetType": "good",
                        "targetId": 5002,
                        "reportType": "虚假宣传",
                        "description": "商品描述与实拍严重不符，且疑似盗用他人图片。",
                        "contactPhone": "13800138000",
                        "status": "PROCESSING",
                        "handleRemark": "已转交平台运营核实图片来源",
                        "attachments": [{"fileId": 703, "url": "http://localhost:9001/api/files/view/beauty-proof.png", "originalName": "beauty-proof.png"}],
                        "handledBy": 9001,
                        "handledAt": "2026-06-04T18:30:00",
                        "createdAt": "2026-06-04T18:00:00",
                    },
                }
            ),
            [("成功返回（非本人举报或记录不存在）", success_lines({"code": 0, "message": "success", "data": None}))],
            [
                "1. 当前详情接口会校验举报人身份，因此只有提交者本人才能查看自己的美妆举报详情。",
                "2. 如果前端打开的举报单不是当前用户提交，接口会返回 data=null，而不是直接报错。",
                "3. 这条接口同样是通用举报能力，但在美妆模块中可以直接复用为“我的美妆举报详情”。",
            ],
        ),
        endpoint(
            "获取美妆个人中心资料接口",
            "美妆模块复用接口",
            "获取美妆个人中心资料",
            "/api/user/profile",
            "GET",
            "application/json",
            "需携带普通用户 Token",
            "获取美妆个人中心展示所需的当前用户资料。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
            ],
            ["GET /api/user/profile", f"Authorization: {USER_TOKEN}"],
            USER_PROFILE_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "userId": 1001,
                        "phone": "13800138000",
                        "nickname": "美妆分享官",
                        "avatarUrl": "http://localhost:9001/api/files/view/avatar.png",
                        "role": "USER",
                        "realNameStatus": "APPROVED",
                        "gender": "女",
                        "contactPhone": "13800138000",
                        "bio": "分享学生党平价好物和宿舍化妆体验",
                    },
                }
            ),
            [("失败返回（未登录）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 美妆个人中心页面当前展示的昵称、手机号等信息来源于通用用户资料接口，而不是独立美妆接口。",
                "2. 前端再结合本地存储中的美妆发布记录与收藏记录，拼出完整的美妆个人中心内容。",
                "3. 因此这条接口主要负责“用户基础档案”，不直接返回美妆商品列表。",
            ],
        ),
        endpoint(
            "更新美妆个人中心资料接口",
            "美妆模块复用接口",
            "更新美妆个人中心资料",
            "/api/user/profile",
            "PUT",
            "application/json",
            "需携带普通用户 Token",
            "更新美妆个人中心中展示的昵称、头像、联系方式等通用资料。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"普通用户令牌，格式 {USER_TOKEN}"),
                param_row("Body", "nickname", "string", "是", "昵称"),
                param_row("Body", "avatarUrl", "string", "否", "头像地址"),
                param_row("Body", "gender", "string", "否", "性别"),
                param_row("Body", "contactPhone", "string", "否", "联系手机号"),
                param_row("Body", "bio", "string", "否", "个人简介"),
            ],
            success_lines(
                {
                    "nickname": "美妆分享官",
                    "avatarUrl": "http://localhost:9001/api/files/view/avatar.png",
                    "gender": "女",
                    "contactPhone": "13800138000",
                    "bio": "分享学生党平价好物和宿舍化妆体验",
                }
            ),
            USER_PROFILE_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "userId": 1001,
                        "phone": "13800138000",
                        "nickname": "美妆分享官",
                        "avatarUrl": "http://localhost:9001/api/files/view/avatar.png",
                        "role": "USER",
                        "realNameStatus": "APPROVED",
                        "gender": "女",
                        "contactPhone": "13800138000",
                        "bio": "分享学生党平价好物和宿舍化妆体验",
                    },
                }
            ),
            [("失败返回（参数不合法）", fail_lines(400, "请求参数不合法"))],
            [
                "1. 美妆个人中心的“编辑个人资料”当前直接跳转通用资料编辑页，因此底层仍是用户资料更新接口。",
                "2. 这意味着美妆模块和其他业务模块共享同一份用户头像、昵称和联系方式。",
                "3. 如果后续要加“美妆达人标签”“常用肤质标签”等专属字段，则需要新增专门的美妆资料接口或扩展用户资料表。",
            ],
        ),
        endpoint(
            "管理员登录接口",
            "后台认证接口",
            "管理员登录",
            "/api/admin/login",
            "POST",
            "application/json",
            "无需登录即可访问",
            "管理员使用手机号和密码登录后台，返回后台使用的 JWT 与用户资料。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Body", "phone", "string", "是", "管理员手机号"),
                param_row("Body", "password", "string", "是", "管理员密码"),
            ],
            success_lines({"phone": "13900009999", "password": "admin123"}),
            AUTH_LOGIN_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "token": "eyJhbGciOiJIUzI1NiJ9.admin",
                        "user": {
                            "userId": 9001,
                            "phone": "13900009999",
                            "nickname": "平台管理员",
                            "avatarUrl": None,
                            "role": "ADMIN",
                            "realNameStatus": "UNSUBMITTED",
                            "gender": None,
                            "contactPhone": None,
                            "bio": None,
                        },
                    },
                }
            ),
            [("失败返回（账号或密码错误）", fail_lines(401, "手机号或密码错误"))],
            [
                "1. 管理员登录与普通用户登录共用同一套认证服务，因此请求体结构完全一致。",
                "2. 后台前端通常把返回 token 保存为 adminToken，并在后续请求中作为管理员凭证使用。",
                "3. 只有角色为 ADMIN 的账号才能访问 /api/admin/** 路径下的管理接口。",
            ],
        ),
        endpoint(
            "获取用户列表接口",
            "后台用户管理接口",
            "获取用户列表",
            "/api/admin/users",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "获取全部用户列表，按创建时间倒序返回。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
            ],
            ["GET /api/admin/users", f"Authorization: {ADMIN_TOKEN}"],
            ADMIN_USER_LIST_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": [
                        {
                            "id": 1001,
                            "phone": "13800138000",
                            "nickname": "小陈",
                            "avatarUrl": None,
                            "role": "USER",
                            "status": "ACTIVE",
                            "createdAt": "2026-06-01T09:00:00",
                        }
                    ],
                }
            ),
            [("失败返回（无管理员权限）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 该接口直接返回用户实体列表，适合后台用户管理首页快速展示。",
                "2. 当前版本不支持分页和筛选，如果数据量增大可在后续迭代中补充分页参数。",
                "3. 若需要查看更完整资料，可继续调用“获取用户详情接口”。",
            ],
        ),
        endpoint(
            "获取用户详情接口",
            "后台用户管理接口",
            "获取用户详情",
            "/api/admin/users/{id}",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "获取指定用户的后台详情视图，含实名状态和举报统计。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Path", "id", "number", "是", "用户 ID"),
            ],
            ["GET /api/admin/users/1001", f"Authorization: {ADMIN_TOKEN}"],
            ADMIN_USER_DETAIL_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "userId": 1001,
                        "phone": "13800138000",
                        "nickname": "小陈",
                        "avatarUrl": None,
                        "role": "USER",
                        "status": "ACTIVE",
                        "realNameStatus": "APPROVED",
                        "gender": "女",
                        "contactPhone": "13800138000",
                        "bio": "可以帮取快递",
                        "reportCount": 1,
                        "createdAt": "2026-06-01T09:00:00",
                    },
                }
            ),
            [("成功返回（用户不存在）", success_lines({"code": 0, "message": "success", "data": None}))],
            [
                "1. 后台详情接口会聚合 user、user_profile、实名记录和举报统计，是用户审核页的主要数据来源。",
                "2. 当前实现对不存在的用户返回 data=null，而不是抛 404 异常。",
                "3. 如果后台需要禁用账号、重置权限等操作，可在此详情视图基础上继续扩展管理动作。",
            ],
        ),
        endpoint(
            "获取实名认证审核列表接口",
            "后台实名认证接口",
            "获取实名认证审核列表",
            "/api/admin/verifications",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "分页获取实名认证记录，默认每页 20 条，按提交时间倒序返回。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Query", "current", "number", "否", "页码，默认 1"),
                param_row("Query", "size", "number", "否", "每页条数，默认 20"),
            ],
            ["GET /api/admin/verifications?current=1&size=20", f"Authorization: {ADMIN_TOKEN}"],
            ADMIN_VERIFICATION_PAGE_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "current": 1,
                        "size": 20,
                        "total": 1,
                        "pages": 1,
                        "records": [
                            {
                                "id": 12,
                                "userId": 1001,
                                "realName": "陈一一",
                                "idCardNo": "440101199901018888",
                                "frontFileId": 301,
                                "backFileId": 302,
                                "status": "PENDING",
                                "rejectReason": None,
                                "reviewedBy": None,
                                "reviewedAt": None,
                                "createdAt": "2026-06-04T10:00:00",
                            }
                        ],
                    },
                }
            ),
            [("失败返回（无管理员权限）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 当前接口名为审核列表，但实际上返回全部实名记录，前端可自行按 status 做筛选。",
                "2. 后端使用 MyBatis-Plus Page 对象分页，records 中存放当前页数据。",
                "3. 列表适合后台审核页与统计页共用，点击单条可继续查看认证详情。",
            ],
        ),
        endpoint(
            "获取实名认证详情接口",
            "后台实名认证接口",
            "获取实名认证详情",
            "/api/admin/verifications/{id}",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "获取指定实名记录详情，便于后台查看证件图和审核状态。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Path", "id", "number", "是", "实名记录 ID"),
            ],
            ["GET /api/admin/verifications/12", f"Authorization: {ADMIN_TOKEN}"],
            VERIFICATION_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 12,
                        "realName": "陈一一",
                        "idCardNo": "440101199901018888",
                        "frontFileId": 301,
                        "frontFileUrl": "http://localhost:9001/api/files/view/id-front.png",
                        "backFileId": 302,
                        "backFileUrl": "http://localhost:9001/api/files/view/id-back.png",
                        "status": "PENDING",
                        "rejectReason": None,
                        "reviewedBy": None,
                        "reviewedAt": None,
                    },
                }
            ),
            [("成功返回（记录不存在）", success_lines({"code": 0, "message": "success", "data": None}))],
            [
                "1. 详情接口会把身份证图片 fileId 转为可直接预览的 URL，方便后台审核人员查看。",
                "2. 当前实现对不存在的实名认证记录返回 data=null。",
                "3. 审核页通常先拉详情，再调用审核接口提交 approve 或 reject 动作。",
            ],
        ),
        endpoint(
            "审核实名认证接口",
            "后台实名认证接口",
            "审核实名认证",
            "/api/admin/verifications/{id}/review",
            "POST",
            "application/json",
            "需携带管理员 Token",
            "对指定实名认证记录执行通过或驳回审核，并向用户发送系统通知。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Path", "id", "number", "是", "实名记录 ID"),
                param_row("Body", "action", "string", "是", "审核动作：approve 或 reject"),
                param_row("Body", "rejectReason", "string", "否", "驳回原因，action=reject 时建议填写"),
            ],
            success_lines({"action": "reject", "rejectReason": "身份证照片不清晰，请重新上传"}),
            [],
            success_lines({"code": 0, "message": "处理成功", "data": None}),
            [
                ("失败返回（实名记录不存在）", fail_lines(400, "实名记录不存在")),
                ("失败返回（参数不合法）", fail_lines(400, "请求参数不合法")),
            ],
            [
                "1. 后台传入 approve 时系统会把状态更新为 APPROVED；否则统一按驳回处理。",
                "2. 审核完成后会记录 reviewedBy、reviewedAt，并给对应用户发送系统通知消息。",
                "3. 同时系统会记录管理员操作日志，便于后续审计和问题追踪。",
            ],
        ),
        endpoint(
            "获取举报列表接口",
            "后台举报管理接口",
            "获取举报列表",
            "/api/admin/reports",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "分页获取全部举报工单，按创建时间倒序返回。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Query", "current", "number", "否", "页码，默认 1"),
                param_row("Query", "size", "number", "否", "每页条数，默认 20"),
            ],
            ["GET /api/admin/reports?current=1&size=20", f"Authorization: {ADMIN_TOKEN}"],
            ADMIN_REPORT_PAGE_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "current": 1,
                        "size": 20,
                        "total": 1,
                        "pages": 1,
                        "records": [
                            {
                                "id": 900,
                                "module": "errand",
                                "targetType": "order",
                                "targetId": 88,
                                "reportType": "未按约履约",
                                "description": "接单后长时间不回复，且未按约送达。",
                                "contactPhone": "13800138000",
                                "status": "PENDING",
                                "submitterId": 1001,
                                "handleRemark": None,
                                "handledBy": None,
                                "handledAt": None,
                                "createdAt": "2026-06-04T13:20:00",
                            }
                        ],
                    },
                }
            ),
            [("失败返回（无管理员权限）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 后台举报列表是平台处理争议、违规和申诉的主要入口。",
                "2. 当前接口返回分页对象，前端应使用 data.records 渲染列表，用 data.total 控制分页器。",
                "3. 若需要查看附件和完整处理信息，可继续调用举报详情接口。",
            ],
        ),
        endpoint(
            "获取举报详情接口",
            "后台举报管理接口",
            "获取举报详情",
            "/api/admin/reports/{id}",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "获取指定举报工单详情，含附件和处理结果。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Path", "id", "number", "是", "举报单 ID"),
            ],
            ["GET /api/admin/reports/900", f"Authorization: {ADMIN_TOKEN}"],
            REPORT_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 900,
                        "module": "errand",
                        "targetType": "order",
                        "targetId": 88,
                        "reportType": "未按约履约",
                        "description": "接单后长时间不回复，且未按约送达。",
                        "contactPhone": "13800138000",
                        "status": "PENDING",
                        "handleRemark": None,
                        "attachments": [{"fileId": 601, "url": "http://localhost:9001/api/files/view/proof.png", "originalName": "proof.png"}],
                        "handledBy": None,
                        "handledAt": None,
                        "createdAt": "2026-06-04T13:20:00",
                    },
                }
            ),
            [("成功返回（记录不存在）", success_lines({"code": 0, "message": "success", "data": None}))],
            [
                "1. 该接口适合后台工单详情页使用，可以完整查看举报内容与证据附件。",
                "2. 当前实现对不存在的举报记录返回 data=null。",
                "3. 管理员处理工单前，通常需要先读取详情再决定是否进入 PROCESSING、RESOLVED 或 REJECTED。",
            ],
        ),
        endpoint(
            "获取跑腿订单管理列表接口",
            "后台跑腿管理接口",
            "获取跑腿订单管理列表",
            "/api/admin/errand/orders",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "分页获取全部跑腿订单，支持按关键字和状态筛选。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Query", "keyword", "string", "否", "关键字，可匹配订单号、地址、发布人/接单人手机号或昵称"),
                param_row("Query", "status", "string", "否", "订单状态过滤值"),
                param_row("Query", "current", "number", "否", "页码，默认 1"),
                param_row("Query", "size", "number", "否", "每页条数，默认 20"),
            ],
            ["GET /api/admin/errand/orders?keyword=小陈&status=PUBLISHED&current=1&size=20", f"Authorization: {ADMIN_TOKEN}"],
            ERRAND_ORDER_PAGE_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "current": 1,
                        "size": 20,
                        "total": 1,
                        "pages": 1,
                        "records": [
                            {
                                "id": 88,
                                "orderNo": "ET20260604120100123",
                                "serviceType": "PICKUP",
                                "pickupAddress": "菜鸟驿站",
                                "deliveryAddress": "2 号宿舍楼 305",
                                "totalFee": 7.0,
                                "status": "PUBLISHED",
                                "publisher": {"nickname": "小陈"},
                                "runner": None,
                                "conversationId": None,
                                "createdAt": "2026-06-04T12:01:00",
                            }
                        ],
                    },
                }
            ),
            [("失败返回（无管理员权限）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 管理端列表会在查询前自动处理超时订单，因此状态筛选能反映真实业务状态。",
                "2. keyword 可同时命中用户手机号、昵称和订单信息，便于后台快速检索。",
                "3. 若需要执行强制取消、恢复公开等动作，可继续调用管理员处理跑腿订单接口。",
            ],
        ),
        endpoint(
            "获取跑腿订单管理详情接口",
            "后台跑腿管理接口",
            "获取跑腿订单管理详情",
            "/api/admin/errand/orders/{id}",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "获取指定跑腿订单详情，管理员可查看全部字段，不受公开状态限制。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Path", "id", "number", "是", "跑腿订单 ID"),
            ],
            ["GET /api/admin/errand/orders/88", f"Authorization: {ADMIN_TOKEN}"],
            ERRAND_ORDER_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 88,
                        "orderNo": "ET20260604120100123",
                        "serviceType": "PICKUP",
                        "pickupAddress": "菜鸟驿站",
                        "deliveryAddress": "2 号宿舍楼 305",
                        "totalFee": 7.0,
                        "status": "DISPUTED",
                        "publisher": {"userId": 1001, "nickname": "小陈", "phone": "13800138000"},
                        "runner": {"userId": 1002, "nickname": "小林", "phone": "13900001111"},
                        "relatedReports": [{"id": 900, "status": "PROCESSING"}],
                        "conversationId": 501,
                        "createdAt": "2026-06-04T12:01:00",
                    },
                }
            ),
            [("失败返回（订单不存在）", fail_lines(400, "跑腿订单不存在"))],
            [
                "1. 管理员详情接口不走普通用户的可见性校验，因此可直接查看已下架、已取消或争议订单。",
                "2. 返回中会带上 relatedReports，方便后台在订单详情页联动查看关联举报。",
                "3. 如果需要直接介入处理订单状态，可在详情页调用管理员动作接口完成处置。",
            ],
        ),
        endpoint(
            "管理员处理跑腿订单接口",
            "后台跑腿管理接口",
            "管理员处理跑腿订单",
            "/api/admin/errand/orders/{id}/action",
            "POST",
            "application/json",
            "需携带管理员 Token",
            "对跑腿订单执行后台干预动作，如强制取消、恢复公开、标记争议、后台确认完成。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Path", "id", "number", "是", "跑腿订单 ID"),
                param_row("Body", "action", "string", "是", "动作：FORCE_CANCEL/RESTORE_PUBLISHED/MARK_DISPUTED/MARK_COMPLETED"),
                param_row("Body", "remark", "string", "否", "处理备注或取消原因"),
            ],
            success_lines({"action": "FORCE_CANCEL", "remark": "联系不上接单人，后台取消"}),
            ERRAND_ORDER_FIELDS,
            success_lines(
                {
                    "code": 0,
                    "message": "success",
                    "data": {
                        "id": 88,
                        "orderNo": "ET20260604120100123",
                        "status": "CANCELLED",
                        "cancelReason": "联系不上接单人，后台取消",
                        "publicVisible": False,
                        "publisher": {"userId": 1001, "nickname": "小陈", "phone": "13800138000"},
                        "runner": {"userId": 1002, "nickname": "小林", "phone": "13900001111"},
                        "conversationId": 501,
                        "createdAt": "2026-06-04T12:01:00",
                    },
                }
            ),
            [
                ("失败返回（订单不存在）", fail_lines(400, "跑腿订单不存在")),
                ("失败返回（动作不支持）", fail_lines(400, "不支持的订单处理动作")),
            ],
            [
                "1. 后台动作接口是平台人工介入跑腿业务的核心能力，可在异常履约、纠纷和误操作场景中兜底。",
                "2. 成功处理后系统会向订单会话推送通知消息，并写入管理员操作日志。",
                "3. RESTORE_PUBLISHED 会清空接单相关字段并重新公开订单，适合误接单或异常恢复场景。",
            ],
        ),
        endpoint(
            "获取跑腿规则配置接口",
            "后台跑腿配置接口",
            "获取跑腿规则配置",
            "/api/admin/errand/rules",
            "GET",
            "application/json",
            "需携带管理员 Token",
            "获取当前后台跑腿费用配置。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
            ],
            ["GET /api/admin/errand/rules", f"Authorization: {ADMIN_TOKEN}"],
            ERRAND_RULE_FIELDS,
            success_lines({"code": 0, "message": "success", "data": {"urgentFee": 2.0, "fragileFee": 3.0}}),
            [("失败返回（无管理员权限）", fail_lines(401, "未登录或 token 已过期"))],
            [
                "1. 该接口与移动端获取规则接口返回结构一致，但定位在后台配置页面。",
                "2. 后台可先读取当前配置，再决定是否调整附加费用。",
                "3. 配置变更后无需重启前端，移动端重新拉取规则即可生效。",
            ],
        ),
        endpoint(
            "更新跑腿规则配置接口",
            "后台跑腿配置接口",
            "更新跑腿规则配置",
            "/api/admin/errand/rules",
            "POST",
            "application/json",
            "需携带管理员 Token",
            "更新加急费和易碎费配置，并返回最新规则。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Body", "urgentFee", "number", "是", "加急附加费，最小值 0"),
                param_row("Body", "fragileFee", "number", "是", "易碎附加费，最小值 0"),
            ],
            success_lines({"urgentFee": 2.5, "fragileFee": 3.5}),
            ERRAND_RULE_FIELDS,
            success_lines({"code": 0, "message": "success", "data": {"urgentFee": 2.5, "fragileFee": 3.5}}),
            [("失败返回（参数不合法）", fail_lines(400, "请求参数不合法"))],
            [
                "1. 当前实现把费用配置保存在运行时配置对象中，因此更新后会立即影响后续新建订单的费用计算。",
                "2. 接口返回更新后的规则对象，后台前端可直接回填到表单中。",
                "3. 若未来需要持久化到数据库，可保持接口路径和请求结构不变，仅调整服务层实现。",
            ],
        ),
        endpoint(
            "处理举报接口",
            "后台举报管理接口",
            "处理举报",
            "/api/admin/reports/{id}/handle",
            "POST",
            "application/json",
            "需携带管理员 Token",
            "更新举报工单状态和处理备注，并给提交人发送通知。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Path", "id", "number", "是", "举报单 ID"),
                param_row("Body", "status", "string", "是", "处理状态：PENDING/PROCESSING/RESOLVED/REJECTED"),
                param_row("Body", "remark", "string", "否", "处理备注"),
            ],
            success_lines({"status": "PROCESSING", "remark": "已联系双方核实，先进入争议处理"}),
            [],
            success_lines({"code": 0, "message": "处理成功", "data": None}),
            [("失败返回（参数不合法）", fail_lines(400, "请求参数不合法"))],
            [
                "1. 处理举报后，系统会记录 handledBy、handledAt，并写入管理员操作日志。",
                "2. 若举报对象是跑腿订单，且状态被更新为 PROCESSING，系统会自动把该订单标记为 DISPUTED。",
                "3. 同时系统会通过 WebSocket 和系统消息双通道通知举报提交人，确保处理进度可感知。",
            ],
        ),
        endpoint(
            "发送系统消息接口",
            "后台消息运营接口",
            "发送系统消息",
            "/api/admin/messages/system/send",
            "POST",
            "application/json",
            "需携带管理员 Token",
            "向指定用户或全部普通用户发送系统通知消息。",
            [
                ["参数位置", "参数名", "数据类型", "是否必填", "参数说明"],
                param_row("Header", "Authorization", "string", "是", f"管理员令牌，格式 {ADMIN_TOKEN}"),
                param_row("Body", "userIds", "number[]", "否", "接收用户 ID 列表；为空时默认发送给全部普通用户"),
                param_row("Body", "content", "string", "是", "系统消息内容"),
            ],
            success_lines({"userIds": [1001, 1002], "content": "系统将于今晚 23:00 进行维护，请提前保存数据。"}),
            [],
            success_lines({"code": 0, "message": "发送成功", "data": None}),
            [("失败返回（参数不合法）", fail_lines(400, "请求参数不合法"))],
            [
                "1. 当 userIds 为空时，后台会自动查询全部 role=USER 的用户并批量发送系统消息。",
                "2. 每个用户都会被写入对应的 SYSTEM 会话中，消息中心可直接查看通知内容。",
                "3. 该接口适合用于平台公告、活动通知、风控提醒等运营场景。",
            ],
        ),
    ]
)


def add_endpoint_block(doc: Document, index: int, item: dict) -> None:
    if item.get("section_header"):
        add_paragraph(doc, item["section_header"], "Heading 2")
    add_paragraph(doc, f"接口{cn_num(index)}：{item['title']}（{item['category']}）", "Heading 2")
    add_paragraph(doc, "1. 基础信息", "Heading 3")
    add_paragraph(doc, f"接口名称：{item['name']}")
    add_paragraph(doc, f"请求地址：{item['path']}")
    add_paragraph(doc, f"请求方式：{item['method']}")
    add_paragraph(doc, f"请求数据类型：{item['content_type']}")
    add_paragraph(doc, f"接口特性：{item['feature']}")
    add_paragraph(doc, f"接口功能：{item['purpose']}")

    add_paragraph(doc, "2. 请求参数（前端传后端）", "Heading 3")
    request_widths = [1.0, 1.35, 1.1, 0.85, 2.5]
    response_widths = [1.6, 1.2, 3.0]
    add_table(doc, item["request_rows"], request_widths)

    add_paragraph(doc, "3. 请求示例", "Heading 3")
    add_lines(doc, item["request_example"])

    add_paragraph(doc, "4. 返回参数（后端传前端）", "Heading 3")
    add_table(doc, item["response_rows"], response_widths)

    add_paragraph(doc, "5. 返回示例", "Heading 3")
    add_paragraph(doc, "成功返回")
    add_lines(doc, item["success_example"])
    for title, lines in item["fail_examples"]:
        add_paragraph(doc, title)
        add_lines(doc, lines)

    add_paragraph(doc, "6. 接口详细解释与协作逻辑", "Heading 3")
    add_lines(doc, item["logic_lines"])


def add_appendix(doc: Document, start_index: int) -> None:
    add_paragraph(doc, f"附录{cn_num(start_index)}：美妆模块接口说明", "Heading 2")
    add_paragraph(doc, "1. 模块现状说明", "Heading 3")
    add_paragraph(doc, "当前版本的美妆模块前端页面位于 mobile/src/pages/beauty/ 目录，核心数据逻辑位于 mobile/src/utils/beautyMock.js。该模块暂未落独立后端 Controller，商品列表、详情、我的发布、我的收藏等数据主要通过前端本地 mock 数据与 localStorage 维护。")
    add_paragraph(doc, "2. 当前复用的正式后端接口", "Heading 3")
    add_paragraph(doc, "美妆模块虽然没有独立后端接口，但实际已经复用了以下正式 HTTP 接口：")
    add_paragraph(doc, "1. POST /api/files/upload/image：用于发布平价好物时上传商品实拍图和订单购买凭证。")
    add_paragraph(doc, "2. POST /api/reports：用于在美妆详情页或列表页对商品发起举报，前端传参通常为 module=beauty、targetType=good、targetId=商品ID。")
    add_paragraph(doc, "3. GET /api/reports 与 GET /api/reports/{id}：用于后续查看本人提交的美妆举报记录与详情。")
    add_paragraph(doc, "4. GET /api/auth/me、GET /api/user/profile、PUT /api/user/profile：用于美妆个人中心展示当前用户资料，以及编辑昵称、联系方式等基础信息。")
    add_paragraph(doc, "3. 当前前端本地数据方法（非 HTTP 接口）", "Heading 3")
    add_paragraph(doc, "以下能力当前属于前端本地方法或缓存读写，不是服务端 HTTP 接口，写文档时需要与正式接口区分开：")
    add_paragraph(doc, "1. getBeautyGoodsList：读取内置 mock 商品和用户本地发布商品，生成美妆好物列表。")
    add_paragraph(doc, "2. getBeautyGoodById：根据商品 ID 获取详情页数据。")
    add_paragraph(doc, "3. createBeautyGood：把用户发布的好物写入本地存储 beauty-user-goods，用于发布页和我的发布页面。")
    add_paragraph(doc, "4. getBeautyFavoriteIds 与 getBeautyFavoriteGoods：基于本地存储 beauty-favorite-ids 维护收藏状态。")
    add_paragraph(doc, "5. getBeautyUserPosts：结合当前登录用户 profile，从本地存储中过滤本人发布的商品。")
    add_paragraph(doc, "4. 美妆模块当前实际交互链路", "Heading 3")
    add_paragraph(doc, "美妆发布页当前的真实链路是：前端先调用图片上传接口上传商品图和购买凭证，再调用前端本地 createBeautyGood 方法把发布内容写入本地缓存；并没有把商品发布到后端数据库。")
    add_paragraph(doc, "美妆举报页当前的真实链路是：前端从美妆商品卡片或详情页跳转举报页，并调用正式举报接口把举报单提交到后端，因此举报能力已经具备真实后端支撑。")
    add_paragraph(doc, "5. 后续后端化建议", "Heading 3")
    add_paragraph(doc, "如果后续要把美妆模块完全纳入正式接口体系，建议补充独立的 beauty Controller，并至少新增商品列表、商品详情、商品发布、我的发布、收藏管理、专题列表等接口。当前这份文档已经先把美妆模块实际复用到的正式接口和本地逻辑边界说明清楚。")

    start_index += 1
    add_paragraph(doc, f"附录{cn_num(start_index)}：接口交互通用规范", "Heading 2")
    add_paragraph(doc, "1. 请求基础约定", "Heading 3")
    add_paragraph(doc, "接口统一服务前缀为 /api，服务默认运行端口为 9001。移动端与后台前端都通过 Authorization 请求头传递 JWT，格式为 Bearer {token}。")
    add_paragraph(doc, "2. 统一返回结构", "Heading 3")
    add_paragraph(doc, "所有接口统一返回 code、message、data 三个顶层字段。成功时 code=0；参数校验失败或业务失败通常返回 400；认证失败通常返回 401；系统异常通常返回 500。")
    add_paragraph(doc, "3. 鉴权规则", "Heading 3")
    add_paragraph(doc, "公开接口只有 /auth/register、/auth/login、/admin/login 以及文件查看、WebSocket 等少量路径。除公开接口外，其余接口都要求登录；/admin/** 路径仅管理员角色可访问。")
    add_paragraph(doc, "4. 文件与图片协作方式", "Heading 3")
    add_paragraph(doc, "实名、举报、头像、跑腿附件等需要图片的业务，建议先调用上传图片接口拿到 fileId，再把 fileId 提交给具体业务接口。")
    add_paragraph(doc, "5. 实时通知能力", "Heading 3")
    add_paragraph(doc, "消息、举报处理和部分跑腿状态变更会通过 WebSocket 额外推送事件，前端可结合接口拉取与实时推送一起使用。")


def main() -> None:
    doc = Document(str(TEMPLATE_PATH))
    clear_document(doc)
    for idx, item in enumerate(endpoints, start=1):
        add_endpoint_block(doc, idx, item)
    add_appendix(doc, len(endpoints) + 1)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
